from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式
import time
#创建一个空白的word文档
#设置1级标题
def set1(str,doc):
    para_heading=doc.add_heading('',level=1)#返回1级标题段落对象，标题也相当于一个段落
    para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
    para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
    para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

    run=para_heading.add_run(str)
    run.font.name='宋体'    #设置为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')#设置为宋体，和上边的一起使用
    run.font.size=Pt(22)#设置1级标题文字的大小为“小四” 为12磅
    run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

def settext(str,doc):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    p.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    p.paragraph_format.line_spacing=1.5  #设置行间距为 1.5倍
    #p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
    p.paragraph_format.first_line_indent=Inches(0.3346457)#相当于小四两个字符的缩进

    p.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    p.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

    r=p.add_run(str)
    r.font.name='宋体'    #设置为宋体
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')#设置为宋体，和上边的一起使用
    r.font.size=Pt(13)  #设置字体大小为12磅 相当于 小四
    r.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

#设置2级标题
def set2(str,doc):
    para_heading=doc.add_heading('',level=2)#返回1级标题段落对象，标题也相当于一个段落
    para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
    para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
    para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

    run=para_heading.add_run(str)
    run.font.name='宋体'    #设置为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')#设置为宋体，和上边的一起使用
    run.font.size=Pt(16)#设置1级标题文字的大小为“小四” 为12磅
    run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

def set3(str,doc):
    para_heading=doc.add_heading('',level=3)#返回1级标题段落对象，标题也相当于一个段落
    para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
    para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
    para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

    run=para_heading.add_run(str)
    run.font.name='宋体'    #设置为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')#设置为宋体，和上边的一起使用
    run.font.size=Pt(14)#设置1级标题文字的大小为“小四” 为12磅
    run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

while True:
    str = input("\n日志信息\n")
    try:
        doc=Document("record.docx")
    except Exception as m:
        pass
        doc=Document()
    time_str = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
    time_str_old = ""
    record=open("record.txt","a+")
    log=open("log.txt","r+")
    logtext = log.readlines()
    try:
        time_str_old = logtext[-1]
    except Exception as m:
        print("warning:没有log文件或文件为空")
        time_str_old = "1995-09-15"
    if time_str[0:4]!=time_str_old[0:4]:
        set1(time_str[0:4]+"年",doc)
        set2(time_str[0:4] + "年" + time_str[5:7] + "月",doc)
        set3(time_str[0:4] + "年" + time_str[5:7] + "月" + time_str[8:10] + "日",doc)
    elif time_str[5:7] != time_str_old[5:7]:
        set2(time_str[0:4]+"年"+time_str[5:7]+"月",doc)
        set3(time_str[0:4] + "年" + time_str[5:7] + "月" + time_str[8:10] + "日",doc)
    elif time_str[8:10] != time_str_old[8:10]:
        set3(time_str[0:4] + "年" + time_str[5:7] + "月" + time_str[8:10]+"日",doc)
    settext(time_str+str,doc)
    record.write(time_str+str+"\n")
    log.write("\n"+time_str)
    log.close()
    record.close()
    doc.save("record.docx")
    print("该次记录时间为"+time_str)


#set2("2021")
#doc.save("测试文件.docx")